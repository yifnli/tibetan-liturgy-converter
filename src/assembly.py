"""
assembly.py

Stage 6 — Assembly.

Takes the resolved verse-unit dicts (produced by combining verse_pipeline,
translation_resolve, and homophonic_resolve outputs) and reflows them into
a python-docx Document object.

Output
------
A python-docx ``Document`` instance.  No file I/O occurs inside this module.
The caller is responsible for writing the document to disk via
``doc.save(path)``.

Layer display order
-------------------
Tibetan is always first (non-configurable).  The remaining layers follow the
``layer_order`` parameter, which defaults to the Rigpe Dorje reference-sample
order: ``["wylie", "french", "english", "mandarin_homophonic",
"mandarin_semantic"]``.

Tibetan rendering
-----------------
- ``UNICODE_TEXT`` path: ``tibetan_unicode`` rendered in ``Jomolhari`` font,
  applied at the run level so the Normal paragraph style is not disturbed.
- ``LEGACY_FONT_TEXT`` path: one inline image per path in
  ``tibetan_image_paths``, all in one paragraph, with a line break between
  images (not after the last).
- Both absent/empty: no Tibetan paragraph is added.

Gap handling
------------
A verse-unit with gaps is **never skipped and never raises**.  Each absent
requested layer becomes an italic placeholder paragraph:

  ``"needs_generation"``    → ``[NEEDS GENERATION — derive from: {source}]``
  ``"novel_needs_review"``  → ``[NOVEL — HUMAN REVIEW REQUIRED — source: {source}]``
  ``"no_source_available"`` → ``[NO SOURCE AVAILABLE]``

Input dicts are never mutated.
"""

from docx import Document

_TIBETAN_FONT = "Jomolhari"
_DEFAULT_LAYER_ORDER: list[str] = [
    "wylie", "french", "english", "mandarin_homophonic", "mandarin_semantic",
]
_VALID_GAP_STATUSES = frozenset({
    "needs_generation", "no_source_available", "novel_needs_review",
})


# ── Function 1 ─────────────────────────────────────────────────────────────

def _make_placeholder_text(gap: dict) -> str:
    """Return the exact placeholder string for a gap record.

    Raises
    ------
    ValueError for any status not in the three valid values.
    """
    status = gap.get("status")
    source = gap.get("source")

    if status == "needs_generation":
        return f"[NEEDS GENERATION \u2014 derive from: {source}]"
    elif status == "novel_needs_review":
        return f"[NOVEL \u2014 HUMAN REVIEW REQUIRED \u2014 source: {source}]"
    elif status == "no_source_available":
        return "[NO SOURCE AVAILABLE]"
    else:
        raise ValueError(
            f"Unknown gap status: {status!r}. "
            f"Valid values: {sorted(_VALID_GAP_STATUSES)}"
        )


# ── Function 2 ─────────────────────────────────────────────────────────────

def _add_tibetan_block(doc: Document, verse_input: dict) -> None:
    """Add the Tibetan paragraph to doc.

    Does nothing if both ``tibetan_unicode`` and ``tibetan_image_paths`` are
    absent / empty.
    """
    tibetan_unicode = verse_input.get("tibetan_unicode")
    tibetan_image_paths = verse_input.get("tibetan_image_paths") or []

    if tibetan_unicode is not None:
        para = doc.add_paragraph()
        run = para.add_run(tibetan_unicode)
        run.font.name = _TIBETAN_FONT

    elif tibetan_image_paths:
        para = doc.add_paragraph()
        for i, path in enumerate(tibetan_image_paths):
            run = para.add_run()
            run.add_picture(path)
            if i < len(tibetan_image_paths) - 1:
                run.add_break()   # line break between images, not after last


# ── Function 3 ─────────────────────────────────────────────────────────────

def _add_layer_paragraph(
    doc: Document,
    layer_name: str,
    value: str | None,
    gap: dict | None,
) -> None:
    """Add exactly one layer paragraph to doc — or nothing if both are absent.

    - ``value`` is not None → plain paragraph with ``value`` as text.
    - ``value`` is None and ``gap`` is not None → italic placeholder paragraph.
    - Both None → no paragraph added (silently ignored; must not raise).
    """
    if value is not None:
        doc.add_paragraph(value)
    elif gap is not None:
        para = doc.add_paragraph()
        run = para.add_run(_make_placeholder_text(gap))
        run.italic = True


# ── Function 4 ─────────────────────────────────────────────────────────────

def add_verse_to_document(
    doc: Document,
    verse_input: dict,
    layer_order: list[str] | None = None,
) -> None:
    """Add all layers for one verse-unit to ``doc`` in the specified order.

    Tibetan is always added first via ``_add_tibetan_block``, regardless of
    what ``layer_order`` contains.  Then each layer in ``layer_order`` is
    added via ``_add_layer_paragraph``.

    Input dicts are not mutated.
    """
    if layer_order is None:
        layer_order = _DEFAULT_LAYER_ORDER

    # Build gap lookup once for O(1) access by layer name.
    gap_lookup: dict[str, dict] = {
        g["layer"]: g for g in verse_input.get("gaps", [])
    }

    _add_tibetan_block(doc, verse_input)

    layers = verse_input.get("layers", {})
    for layer_name in layer_order:
        _add_layer_paragraph(
            doc, layer_name,
            layers.get(layer_name),
            gap_lookup.get(layer_name),
        )


# ── Function 5 ─────────────────────────────────────────────────────────────

def build_document(
    verse_inputs: list[dict],
    layer_order: list[str] | None = None,
) -> Document:
    """Create a new ``Document``, add every verse-unit, return the document.

    The caller writes to disk via ``doc.save(path)``.  Input dicts are not
    mutated.
    """
    if layer_order is None:
        layer_order = _DEFAULT_LAYER_ORDER

    doc = Document()
    for verse_input in verse_inputs:
        add_verse_to_document(doc, verse_input, layer_order)
    return doc
