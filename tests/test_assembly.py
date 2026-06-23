# tests/test_assembly.py
#
# Tests for src/assembly.py — Stage 6 assembly.
# Tests added cumulatively, one function at a time.

import sys
import struct
import zlib
import copy

import pytest
from docx import Document

sys.path.insert(0, "src")

from src.assembly import (
    _make_placeholder_text,
    _add_tibetan_block,
    _add_layer_paragraph,
    add_verse_to_document,
    build_document,
)


# ── Shared helper: minimal valid 1×1 white RGB PNG ────────────────────────

def _write_minimal_png(path: str) -> None:
    """Write a minimal valid 1×1 white RGB PNG using stdlib only (no Pillow)."""
    def chunk(tag: bytes, data: bytes) -> bytes:
        crc = zlib.crc32(tag + data) & 0xFFFFFFFF
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", crc)

    # IHDR: width=1, height=1, bit_depth=8, color_type=2 (RGB), rest=0
    ihdr = chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
    # IDAT: filter byte 0 (None) + R G B = white pixel
    idat = chunk(b"IDAT", zlib.compress(b"\x00\xff\xff\xff"))
    iend = chunk(b"IEND", b"")
    with open(path, "wb") as f:
        f.write(b"\x89PNG\r\n\x1a\n" + ihdr + idat + iend)


# ── Shared live-data constants (from explore_assembly_inputs.py output) ───
#
# Strings copied verbatim — do not retype from memory.

_TIBETAN_UNICODE_RD = (
    "༈ རྡོ་རྗེ་འཆང་ཆེན་ཏཻ་ལོ་ནཱ་རོ་དང༌། "
    "།མར་པ་མི་ལ་ཆོས་རྗེ་སྒམ་པོ་པ། །"
)

_ASSEMBLY_INPUT_RD = {
    "tibetan_unicode":     _TIBETAN_UNICODE_RD,
    "tibetan_image_paths": [],
    "layers": {
        "wylie":               "DOR JE CHANG CHEN TI LO NA RO DANG  MAR PA MI LA CHÖ JE GAM PO PA",
        "french":              "Grand Vajradhara, Tilopa, Naropa, Marpa, Milarépa, Seigneur-du-Dharma Gampopa,",
        "english":             "Great Vajradhara, Tilopa, Naropa, Marpa, Milarepa, Lord of Dharma Gampopa,",
        "mandarin_homophonic": "多杰 羌千 帝洛 那诺倘 玛巴 ⽶拉 却戒 冈波巴",
        "mandarin_semantic":   "⾦刚总持帝洛那洛巴，马巴密勒法王冈波巴，",
    },
    "gaps":      [],
    "verse_ref": {"page_index": 2, "verse_index": 0},
    "doc_format": "UNICODE_TEXT",
}


def _make_mb_assembly_input(png_path: str) -> dict:
    """Build MB-style assembly input with a synthetic 1×1 PNG image path."""
    return {
        "tibetan_unicode":     None,
        "tibetan_image_paths": [png_path],
        "layers": {
            "wylie":               "DOR JE CHANG CHEN TE LO NA RO DANG",
            "french":              "G reat Vajradhara, Tilopa, Naropa,",
            "english":             "Grand Vajradhara, Tilopa, Naropa",
            "mandarin_homophonic": None,
            "mandarin_semantic":   None,
        },
        "gaps": [
            {
                "layer": "mandarin_semantic",
                "source": "english",
                "status": "needs_generation",
                "verse_ref": {"page_index": 1, "verse_index": 0},
            },
            {
                "layer": "mandarin_homophonic",
                "source": "wylie",
                "status": "novel_needs_review",
                "verse_ref": {"page_index": 1, "verse_index": 0},
            },
        ],
        "verse_ref": {"page_index": 1, "verse_index": 0},
        "doc_format": "LEGACY_FONT_TEXT",
    }


# ── Tests 1–4: _make_placeholder_text ────────────────────────────────────

def test_placeholder_needs_generation_exact_string():
    gap = {"status": "needs_generation", "source": "english",
           "layer": "mandarin_semantic", "verse_ref": {}}
    assert _make_placeholder_text(gap) == "[NEEDS GENERATION \u2014 derive from: english]"


def test_placeholder_novel_needs_review_exact_string():
    gap = {"status": "novel_needs_review", "source": "wylie",
           "layer": "mandarin_homophonic", "verse_ref": {}}
    assert _make_placeholder_text(gap) == "[NOVEL \u2014 HUMAN REVIEW REQUIRED \u2014 source: wylie]"


def test_placeholder_no_source_available_exact_string():
    gap = {"status": "no_source_available", "source": None,
           "layer": "mandarin_homophonic", "verse_ref": {}}
    assert _make_placeholder_text(gap) == "[NO SOURCE AVAILABLE]"


def test_placeholder_unknown_status_raises_value_error():
    gap = {"status": "invalid_status", "source": "english",
           "layer": "french", "verse_ref": {}}
    with pytest.raises(ValueError, match="Unknown gap status"):
        _make_placeholder_text(gap)


# ── Tests 5–7: _add_tibetan_block ─────────────────────────────────────────────

def test_tibetan_block_unicode_run_text_and_font():
    verse_input = {"tibetan_unicode": _TIBETAN_UNICODE_RD, "tibetan_image_paths": []}
    doc = Document()
    initial = len(doc.paragraphs)
    _add_tibetan_block(doc, verse_input)
    assert len(doc.paragraphs) == initial + 1
    run = doc.paragraphs[-1].runs[0]
    assert run.text == _TIBETAN_UNICODE_RD
    assert run.font.name == "Jomolhari"


def test_tibetan_block_image_adds_correct_run_count(tmp_path):
    # 1 image → 1 paragraph with 1 run; no exception raised.
    png = tmp_path / "tibetan.png"
    _write_minimal_png(str(png))
    verse_input = {"tibetan_unicode": None, "tibetan_image_paths": [str(png)]}
    doc = Document()
    initial = len(doc.paragraphs)
    _add_tibetan_block(doc, verse_input)
    assert len(doc.paragraphs) == initial + 1
    assert len(doc.paragraphs[-1].runs) == 1


def test_tibetan_block_both_absent_no_paragraph_added():
    verse_input = {"tibetan_unicode": None, "tibetan_image_paths": []}
    doc = Document()
    initial = len(doc.paragraphs)
    _add_tibetan_block(doc, verse_input)
    assert len(doc.paragraphs) == initial


# ── Tests 8–10: _add_layer_paragraph ───────────────────────────────────────────

def test_layer_paragraph_value_present_plain_text():
    doc = Document()
    initial = len(doc.paragraphs)
    _add_layer_paragraph(doc, "english", "Some English text", None)
    assert len(doc.paragraphs) == initial + 1
    para = doc.paragraphs[-1]
    assert para.text == "Some English text"
    assert para.runs[0].italic is not True  # not explicitly italic


def test_layer_paragraph_gap_present_italic_placeholder():
    gap = {"status": "needs_generation", "source": "english",
           "layer": "mandarin_semantic", "verse_ref": {}}
    doc = Document()
    initial = len(doc.paragraphs)
    _add_layer_paragraph(doc, "mandarin_semantic", None, gap)
    assert len(doc.paragraphs) == initial + 1
    para = doc.paragraphs[-1]
    assert para.text == "[NEEDS GENERATION — derive from: english]"
    assert para.runs[0].italic is True


def test_layer_paragraph_both_none_no_paragraph_added():
    doc = Document()
    initial = len(doc.paragraphs)
    _add_layer_paragraph(doc, "french", None, None)
    assert len(doc.paragraphs) == initial


# ── Tests 11–14: add_verse_to_document ──────────────────────────────────────

def test_add_verse_case1_rd_6_paragraphs_no_italic():
    # RD verse: 1 Tibetan + 5 layers (all present) = 6 paragraphs, 0 italic.
    doc = Document()
    initial = len(doc.paragraphs)
    add_verse_to_document(doc, _ASSEMBLY_INPUT_RD)
    added = doc.paragraphs[initial:]
    assert len(added) == 6
    italic_count = sum(1 for p in added if p.runs and p.runs[0].italic is True)
    assert italic_count == 0


def test_add_verse_case2_mb_6_paragraphs_2_italic(tmp_path):
    # MB verse: 1 Tibetan image + 3 text + 2 placeholders = 6 paragraphs, 2 italic.
    png = tmp_path / "tibetan.png"
    _write_minimal_png(str(png))
    mb_input = _make_mb_assembly_input(str(png))
    doc = Document()
    initial = len(doc.paragraphs)
    add_verse_to_document(doc, mb_input)
    added = doc.paragraphs[initial:]
    assert len(added) == 6
    italic_count = sum(1 for p in added if p.runs and p.runs[0].italic is True)
    assert italic_count == 2


def test_add_verse_default_order_paragraph_texts():
    # Default layer_order: wylie → french → english → mandarin_homophonic → mandarin_semantic.
    doc = Document()
    initial = len(doc.paragraphs)
    add_verse_to_document(doc, _ASSEMBLY_INPUT_RD)
    added = doc.paragraphs[initial:]
    # added[0] = Tibetan text, added[1..5] = layers in order
    # Reference the shared constants to avoid re-typing Unicode characters.
    assert added[1].text == _ASSEMBLY_INPUT_RD["layers"]["wylie"]
    assert added[2].text == _ASSEMBLY_INPUT_RD["layers"]["french"]
    assert added[3].text == _ASSEMBLY_INPUT_RD["layers"]["english"]
    assert added[4].text == _ASSEMBLY_INPUT_RD["layers"]["mandarin_homophonic"]
    assert added[5].text == _ASSEMBLY_INPUT_RD["layers"]["mandarin_semantic"]


def test_add_verse_custom_layer_order_one_layer():
    # Custom layer_order=["english"] → exactly 2 paragraphs: Tibetan + English.
    doc = Document()
    initial = len(doc.paragraphs)
    add_verse_to_document(doc, _ASSEMBLY_INPUT_RD, layer_order=["english"])
    added = doc.paragraphs[initial:]
    assert len(added) == 2
    assert added[1].text == "Great Vajradhara, Tilopa, Naropa, Marpa, Milarepa, Lord of Dharma Gampopa,"


# ── Tests 15–18: build_document ───────────────────────────────────────────────

def test_build_document_empty_list_zero_paragraphs():
    # python-docx 1.2.0: Document() starts with 0 paragraphs (no default empty paragraph).
    doc = build_document([])
    assert len(doc.paragraphs) == 0


def test_build_document_returns_document_instance():
    from docx.document import Document as DocxDocument  # factory fn ≠ class
    doc = build_document([])
    assert isinstance(doc, DocxDocument)


def test_build_document_two_verses_paragraph_count():
    # 0 default + 6 (RD verse 1) + 6 (RD verse 2) = 12 paragraphs.
    doc = build_document([_ASSEMBLY_INPUT_RD, _ASSEMBLY_INPUT_RD])
    assert len(doc.paragraphs) == 12


def test_build_document_does_not_mutate_input():
    original = copy.deepcopy([_ASSEMBLY_INPUT_RD])
    inputs = copy.deepcopy([_ASSEMBLY_INPUT_RD])
    build_document(inputs)
    assert inputs == original
